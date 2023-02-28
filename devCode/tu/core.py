# import numpy as np
import datetime
import os
import matplotlib.pyplot as plt
from matplotlib import cm, rcParams
from mpl_toolkits.mplot3d import Axes3D
from numpy.core.function_base import linspace
import pandas as pd
from rich import console
from rich.table import Column, Table
import numpy as np
from scipy import linalg
import math
import pandas as pd
import sys
from docx import Document  # 用来建立一个word对象
from docx.shared import Cm, Pt  # 用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  # 设置字体
from docx.shared import RGBColor  # 设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # 设置对其方式

# from itertools import combinations
# from scipy.special import comb

from rich.console import Console


rcParams["xtick.direction"] = "in"
rcParams["ytick.direction"] = "in"


def mkdir(path):
    folder = os.path.exists(path)

    if not folder:  # 判断是否存在文件夹如果不存在则创建为文件夹
        os.makedirs(path)  # makedirs 创建文件时如果路径不存在会创建这个路径
        print("---  创建算例文件夹于" + path + "  ---")
        print("---  初始化已完成  ---")

    else:
        print("---  文件夹已存在！  ---")


curr_time = datetime.datetime.now()
time_str = datetime.datetime.strftime(curr_time, "%Y-%m-%d-%H-%M-%S")

input_choose = input("是否导入参数（y/n）")

if input_choose == "n":
    name = input("算例文件夹备注：")  # 手动输入
    file = ".\\result\\" + time_str + "-" + name
    mkdir(file)

    # 创建文档
    doc = Document()

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.bold = True

    # 标题设置
    para_heading = doc.add_heading("", level=0)
    para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置为左对齐
    para_heading.paragraph_format.space_before = Pt(0)  # 设置段前 0 磅
    para_heading.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
    para_heading.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5
    para_heading.paragraph_format.left_indent = Inches(0)  # 设置左缩进 1英寸
    para_heading.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0.5 英寸

    run = para_heading.add_run("分析报告_" + name + "_" + time_str)
    run.font.name = u"宋体"  # 设置为宋体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")  # 设置为宋体，和上边的一起使用
    run.font.size = Pt(16)  # 设置1级标题文字的大小为“小四” 为18磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    nn = 0
    while 1:
        nn += 1
        # 设置一级标题
        para_heading = doc.add_heading("", level=1)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run("案例" + str(nn))
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        console = Console()
        """输入参数"""
        console.print("请输入已知参数：", style="green")
        S1 = float(console.input("最小水平主应力（MPa, 参考值：10）："))
        n = float(console.input("幂律系数（无因次，参考值：2）："))
        nl = 2 * n / (2 + n)  # 缝长比例系数
        frac_num = int(console.input("裂缝条数（参考：2，3，4，5）："))
        frac_spac = float(console.input("相邻簇间距（m，参考值：20）："))
        Hw = float(console.input("缝高（m，参考值：20）："))
        Hp = float(console.input("滤失高度（m，参考值：20）："))
        Q = float(console.input("排量（方/s，参考值：0.2）："))
        tp = float(console.input("泵注时间（s，参考值：20）："))
        E = float(console.input("杨氏模量（MPa，参考值：30000）："))
        v = float(console.input("泊松比（无因次，参考值：0.25）："))
        ISIP = float(console.input("瞬时停泵压力（MPa，参考值：15）："))
        P_i = float(console.input("滤失过程i点压力（MPa，参考值：大于等于最小水平主应力,10.1228）："))
        t_i = float(console.input("滤失过程i点闭合时间（s，参考值：小于等于裂缝完全闭合时间,11236.1）："))
        P_j = float(console.input("滤失过程j点压力（MPa，参考值：大于等于最小水平主应力,11.7964）："))
        t_j = float(console.input("滤失过程j点闭合时间（s，参考值：小于等于裂缝完全闭合时间,3016.13）："))
        m = float(console.input("裂缝延伸系数（无因次，参考值：1.5）："))
        frac_type = console.input("裂缝模型（PKN/KGD）：")
        # eff = 25.450  # 压裂液效率

        table = Table(title="输入参数汇总", show_header=True, header_style="bold magenta")
        table.add_column("参数", style="dim", justify="center")
        table.add_column("数值", style="dim", justify="center")
        table.add_row("最小水平主应力,MPa", str(S1))
        table.add_row("幂律系数", str(n))
        table.add_row("缝长比例系数", str(nl))
        table.add_row("裂缝条数", str(frac_num))
        table.add_row("相邻簇间距,m", str(frac_spac))
        table.add_row("缝高,m", str(Hw))
        table.add_row("滤失高度,m", str(Hw))
        table.add_row("排量,方/s", str(Q))
        table.add_row("泵注时间,s", str(tp))
        table.add_row("杨氏模量,MPa", str(E))
        table.add_row("泊松比", str(v))
        table.add_row("瞬时停泵压力,MPa", str(ISIP))
        table.add_row("滤失过程i点压力,MPa", str(P_i))
        table.add_row("滤失过程i点闭合时间,s", str(t_i))
        table.add_row("滤失过程j点压力,MPa", str(P_j))
        table.add_row("滤失过程j点闭合时间,s", str(t_j))
        table.add_row("裂缝延伸系数", str(m))
        table.add_row("裂缝模型", frac_type)
        console.print(table)

        data = [
            ["最小水平主应力,MPa", str(S1)],
            ["幂律系数", str(n)],
            ["缝长比例系数", str(nl)],
            ["裂缝条数", str(frac_num)],
            ["相邻簇间距,m", str(frac_spac)],
            ["缝高,m", str(Hw)],
            ["滤失高度,m", str(Hw)],
            ["排量,方/s", str(Q)],
            ["泵注时间,s", str(tp)],
            ["杨氏模量,MPa", str(E)],
            ["泊松比", str(v)],
            ["瞬时停泵压力,MPa", str(ISIP)],
            ["滤失过程i点压力,MPa", str(P_i)],
            ["滤失过程i点闭合时间,s", str(t_i)],
            ["滤失过程j点压力,MPa", str(P_j)],
            ["滤失过程j点闭合时间,s", str(t_j)],
            ["裂缝延伸系数", str(m)],
            ["裂缝模型", frac_type],
        ]
        df = pd.DataFrame(data, columns=["输入参数", "数值"])

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".1 输入参数")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 创建输入参数表
        table = doc.add_table(df.shape[0] + 1, df.shape[1], style="Light Shading")
        # 插入表头
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        # 插入数据
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        # 设置整个表格字体属性
        table.style.font.size = Pt(12)
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        os.system("pause")

        """模型计算"""
        if frac_type == "PKN":
            beta = math.pi / 4.0
            print(frac_type)
        else:
            beta = 1
        print(beta)

        # 确定簇间距和诱导应力系数矩阵
        d = np.zeros((frac_num, frac_num))  # 簇间距
        phi = np.zeros((frac_num, frac_num))  # 诱导应力系数矩阵
        for i in range(frac_num):
            for j in range(frac_num):
                if i != j:
                    d[i][j] = frac_spac * abs(i - j)
                    phi[i][j] = 1 - math.pow(1 + math.pow(Hw / 2 / d[i][j], 2), -1.5)
        console.print("簇间距矩阵：\n", d, style="green")
        console.print("诱导应力系数矩阵：\n", phi, style="green")

        for i in range(frac_num):
            tmp = 0
            for j in range(frac_num):
                tmp += phi[i][j]
            # console.print(tmp, style="green")
            if tmp >= 1:
                console.print("第%d条裂缝未张开" % (i + 1), style="green")
                # for k in range(frac_num):
                #     phi[k][i] = 0

        console.print("修正后的诱导应力系数矩阵：\n", phi, style="green")

        a = np.zeros((frac_num))  # 缝长系数矩阵
        Pnet = np.zeros((frac_num))  # 停泵时刻各缝净压力矩阵
        for i in range(frac_num):
            tmp = 0
            for j in range(frac_num):
                tmp += phi[i][j]
            a[i] = 1 - tmp
            if a[i] < 0:
                a[i] = 0
            Pnet[i] = a[i] * (ISIP - S1)

        console.print("缝长系数矩阵:\n", a, style="green")
        console.print("停泵时刻各缝净压力矩阵:\n", Pnet, style="green")

        E1 = E / (1.0 - pow(v, 2))  # 平面弹性模量
        ti = (t_i - tp) * 1.0 / tp  # 无因次时间
        gi = (4.0 / math.pi) * (
            (1.0 + ti) * math.asin(math.pow((1.0 + ti), -0.5))
            + math.pow(ti, 0.5)
            - math.pi / 2
        )
        tj = (t_j - tp) * 1.0 / tp  # 无因次时间
        gj = (4.0 / math.pi) * (
            (1.0 + tj) * math.asin(math.pow((1.0 + tj), -0.5))
            + math.pow(tj, 0.5)
            - math.pi / 2
        )
        console.print("第i点无因次闭合时间:", gi)
        console.print("第j点无因次闭合时间:", gj)
        b = (16 * beta * math.pow(Hw, 2) / (5 * E1)) * (
            (ISIP - S1)
            - (P_j - P_i)
            / (gj - gi)
            * (np.sqrt(math.pi) * m * math.gamma(m))
            / ((m + 0.5) * math.gamma(m + 0.5))
        )  # 常数分母项
        c = Q * tp  # 常数分子项
        console.print("b, c:", b, c, style="green")
        console.print("常数项:", (c / b), style="green")

        # 联立线性方程组
        # 联立后的缝长系数矩阵
        a_sum = np.zeros((frac_num, frac_num))  # 联立后的缝长系数矩阵
        for i in range(frac_num - 1):
            a_sum[i][i] = 1
            if (a[i] > 0) and (a[i + 1] > 0):
                a_sum[i][i + 1] = 0 - math.pow(a[i] / a[i + 1], nl)

        for i in range(frac_num):
            a_sum[frac_num - 1][i] = a[i]
        console.print("联立后的缝长系数矩阵:\n", a_sum, style="green")

        b_sum = np.zeros((frac_num))  # 联立后的常数项矩阵
        for i in range(frac_num - 1):
            b_sum[i] = 0
        b_sum[frac_num - 1] = c / b
        console.print("联立后的常数项矩阵:\n", b_sum, style="green")

        Lp = linalg.solve(a_sum, b_sum)
        L = Lp * 2

        Wmax = 2 * Hw * Pnet / E1
        C = 0 - 4 * beta * math.pow(Hw, 2) * a / (5 * E1 * Hp * np.sqrt(tp)) * (
            P_j - P_i
        ) / (gj - gi)

        leak_pump = 0  # 泵注过程的总滤失体积
        for i in range(frac_num):
            leak_pump += (
                -16
                * beta
                * math.pow(Hw, 2)
                * (L[i] / 2)
                / (5 * E1)
                * a[i]
                * (P_j - P_i)
                / (gj - gi)
                * np.sqrt(math.pi)
                * m
                * math.gamma(m)
                / ((m + 0.5) * math.gamma(m + 0.5))
            )
            # (4 * np.sqrt(math.pi) * m * math.gamma(m) * C[i] * Hp * L[i] * np.sqrt(tp))/ ((m + 0.5) * math.gamma(m + 0.5))
        leak_shut = 0  # 停泵过程的总滤失体积
        for i in range(frac_num):
            leak_shut += 16 * beta * math.pow(Hw, 2) * (L[i] / 2) / (5 * E1) * Pnet[i]
        V_frac = 0  # 停泵时刻裂缝总体积
        for i in range(frac_num):
            V_frac += 16 * beta * math.pow(Hw, 2) * (L[i] / 2) / (5 * E1) * Pnet[i]
        eff = (Q * tp - leak_pump) / (Q * tp) * 100

        console.print("缝长:\n", L, style="green")
        console.print("泵注过程的总滤失体积:%.3f" % leak_pump, style="green")
        console.print("停泵过程的总滤失体积:%.3f" % leak_shut, style="green")
        console.print("停泵时刻的裂缝总体积:%.3f" % V_frac, style="green")
        console.print("总泵入体积:%.3f" % (Q * tp), style="green")
        console.print("压裂液效率:%.3f%%" % eff, style="green")
        console.print("最大缝宽:\n", Wmax, style="green")
        console.print("滤失系数:\n", C, style="green")

        # 计算滤失面积
        # s = L * 2 * Hw
        # console.print('滤失面积:\n', s, style="green")
        console.print("计算完毕！", style="green")

        table = Table(title="计算参数汇总", show_header=True, header_style="bold magenta")
        table.add_column("参数", style="dim", justify="center")
        table.add_column("数值", style="dim", justify="center")
        table.add_row("缝长", str(L))
        table.add_row("泵注过程的总滤失体积", str(leak_pump))
        table.add_row("停泵过程的总滤失体积/改造体积", str(leak_shut))
        table.add_row("总泵入体积", str(Q * tp))
        table.add_row("压裂液效率", str(eff))
        table.add_row("最大缝宽", str(Wmax))
        table.add_row("滤失系数", str(C))
        # table.add_row('滤失面积', str(s))
        console.print(table)

        data = [
            ["缝长", str(L)],
            ["泵注过程的总滤失体积", str(leak_pump)],
            ["停泵过程的总滤失体积/改造体积", str(leak_shut)],
            ["总泵入体积", str(Q * tp)],
            ["压裂液效率", str(eff)],
            ["最大缝宽", str(Wmax)],
            ["滤失系数", str(C)],
        ]
        df = pd.DataFrame(data, columns=["输出参数", "数值"])

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".2 输出参数")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 创建输入参数表
        table = doc.add_table(df.shape[0] + 1, df.shape[1], style="Light Shading")
        # 插入表头
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        # 插入数据
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        # 设置整个表格字体属性
        table.style.font.size = Pt(12)
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        """可视化"""
        choose = console.input("是否可视化（y/n）：")

        fig = plt.figure()
        ax = Axes3D(fig, auto_add_to_figure=False)
        fig.add_axes(ax)
        for i in range(frac_num):
            Y = np.arange(0, int(L[i] / 2) + 1, 1)
            Z = np.arange(-Hw, Hw + 1, 1)
            Y, Z = np.meshgrid(Y, Z)
            w = Wmax[i] * np.sqrt(np.sqrt(1 - Y / int(L[i] / 2)))
            X = w * np.sqrt(1 - Z * Z / (Hw * Hw))
            # console.print('x:', X)
            # console.print('y:', Y)
            # console.print('w:', w)
            # console.print('z:', Z)
            ax.plot_surface(
                X + i * max(Wmax) * 3, Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, -Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, -Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, -Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, -Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
        ax.set_xlabel("Fracture Width", rotation=-15)
        ax.set_ylabel("Fracture Length", rotation=50)
        ax.set_zlabel("Fracture Height", rotation=90)

        # 绘制井筒
        # u = np.linspace(0,2*np.pi,50)
        # h = np.linspace(-max(Wmax) * 2,max(Wmax) * 6)
        # y = np.outer(np.sin(u),np.ones(len(h)))
        # z = np.outer(np.cos(u),np.ones(len(h)))
        # x = np.outer(np.ones(len(u)),h)
        # ax.plot_surface(x, y, z, cmap=cm.viridis)

        curr_time = datetime.datetime.now()
        time_str = datetime.datetime.strftime(curr_time, "%Y-%m-%d-%H-%M-%S")
        plt.savefig(file + "\\3D_" + time_str + ".png")
        if choose == "y":
            plt.show()

        fig, ax = plt.subplots()
        for i in range(frac_num):
            Y = np.arange(0, int(L[i] / 2) + 1, 0.001)
            w = Wmax[i] * np.sqrt(np.sqrt(1 - Y / int(L[i] / 2)))  # 缝宽作为颜色
            X = np.full((len(Y)), i * frac_spac)

            ax.scatter(X, Y, marker=",", c=-w, cmap="RdYlGn")
            ax.scatter(X, -Y, marker=",", c=-w, cmap="RdYlGn")
        # 设置坐标轴范围
        plt.xticks(linspace(-frac_spac / 2, (frac_num - 1) * frac_spac + frac_spac / 2, 10))
        plt.yticks(linspace(-max(Y) - frac_spac / 2, max(Y) + frac_spac / 2, 10))
        plt.savefig(file + "\\2D_" + time_str + ".png")
        if choose == "y":
            plt.show()

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".3 结果可视化")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(
            file + "\\3D_" + time_str + ".png", width=Cm(12), height=Cm(8)
        )
        paragraph.add_run().add_picture(
            file + "\\2D_" + time_str + ".png", width=Cm(12), height=Cm(8)
        )

        choose = console.input("重新计算（y/n）：")
        if choose == "n":
            break
        doc.add_section()
        
    doc.save(file + "\\分析报告_" + name + ".docx")
    console.print("分析报告已生成！", style="red")


else:
    name = input("导入文件名：")  # 自动输入
    file = ".\\result\\" + time_str + "-" + name
    mkdir(file)

    # 创建文档
    doc = Document()

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.bold = True

    # 标题设置
    para_heading = doc.add_heading("", level=0)
    para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置为左对齐
    para_heading.paragraph_format.space_before = Pt(0)  # 设置段前 0 磅
    para_heading.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
    para_heading.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5
    para_heading.paragraph_format.left_indent = Inches(0)  # 设置左缩进 1英寸
    para_heading.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0.5 英寸

    run = para_heading.add_run("分析报告_" + name + "_" + time_str)
    run.font.name = u"宋体"  # 设置为宋体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")  # 设置为宋体，和上边的一起使用
    run.font.size = Pt(16)  # 设置1级标题文字的大小为“小四” 为18磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    data_input = pd.ExcelFile(name + ".xlsx")
    for nn in range(1, int(len(data_input.sheet_names)) + 1):
        # 设置一级标题
        para_heading = doc.add_heading("", level=1)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run("案例" + str(nn))
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        console = Console()
        """输入参数"""
        data = data_input.parse(
            sheet_name=data_input.sheet_names[nn - 1], names=["输入参数", "数值"]
        )

        S1 = float(data.values[0, 1]) # 最小水平主应力（MPa, 参考值：10）
        n = float(data.values[1, 1]) # 幂律系数（无因次，参考值：2）
        nl = float(data.values[2, 1]) # 缝长比例系数, nl = 2 * n / (2 + n) 
        frac_num = int(data.values[3, 1]) # 裂缝条数（参考：2，3，4，5）
        frac_spac = float(data.values[4, 1]) # 相邻簇间距（m，参考值：20）
        Hw = float(data.values[5, 1]) # 缝高（m，参考值：20）
        Hp = float(data.values[6, 1]) # 滤失高度（m，参考值：20）
        Q = float(data.values[7, 1]) # 排量（方/s，参考值：0.2）
        tp = float(data.values[8, 1]) # 泵注时间（s，参考值：20）
        E = float(data.values[9, 1]) # 杨氏模量（MPa，参考值：30000）
        v = float(data.values[10, 1]) # 泊松比（无因次，参考值：0.25）
        ISIP = float(data.values[11, 1]) # 瞬时停泵压力（MPa，参考值：15）
        P_i = float(data.values[12, 1]) # 滤失过程i点压力（MPa，参考值：大于等于最小水平主应力,10.1228）
        t_i = float(data.values[13, 1]) # 滤失过程i点闭合时间（s，参考值：小于等于裂缝完全闭合时间,11236.1）
        P_j = float(data.values[14, 1]) # 滤失过程j点压力（MPa，参考值：大于等于最小水平主应力,11.7964）
        t_j = float(data.values[15, 1]) # 滤失过程j点闭合时间（s，参考值：小于等于裂缝完全闭合时间,3016.13）
        m = float(data.values[16, 1]) # 裂缝延伸系数（无因次，参考值：1.5）
        frac_type = str(data.values[17, 1]) # 裂缝模型（PKN/KGD）
        # eff = 25.450  # 压裂液效率
        eff_c = 1  # 泵注效率控制系数

        table = Table(title="输入参数汇总", show_header=True, header_style="bold magenta")
        table.add_column("参数", style="dim", justify="center")
        table.add_column("数值", style="dim", justify="center")
        table.add_row("最小水平主应力,MPa", str(S1))
        table.add_row("幂律系数", str(n))
        table.add_row("缝长比例系数", str(nl))
        table.add_row("裂缝条数", str(frac_num))
        table.add_row("相邻簇间距,m", str(frac_spac))
        table.add_row("缝高,m", str(Hw))
        table.add_row("滤失高度,m", str(Hw))
        table.add_row("排量,方/s", str(Q))
        table.add_row("泵注时间,s", str(tp))
        table.add_row("杨氏模量,MPa", str(E))
        table.add_row("泊松比", str(v))
        table.add_row("瞬时停泵压力,MPa", str(ISIP))
        table.add_row("滤失过程i点压力,MPa", str(P_i))
        table.add_row("滤失过程i点闭合时间,s", str(t_i))
        table.add_row("滤失过程j点压力,MPa", str(P_j))
        table.add_row("滤失过程j点闭合时间,s", str(t_j))
        table.add_row("裂缝延伸系数", str(m))
        table.add_row("裂缝模型", frac_type)
        console.print(table)

        data = [
            ["最小水平主应力,MPa", str(S1)],
            ["幂律系数", str(n)],
            ["缝长比例系数", str(nl)],
            ["裂缝条数", str(frac_num)],
            ["相邻簇间距,m", str(frac_spac)],
            ["缝高,m", str(Hw)],
            ["滤失高度,m", str(Hw)],
            ["排量,方/s", str(Q)],
            ["泵注时间,s", str(tp)],
            ["杨氏模量,MPa", str(E)],
            ["泊松比", str(v)],
            ["瞬时停泵压力,MPa", str(ISIP)],
            ["滤失过程i点压力,MPa", str(P_i)],
            ["滤失过程i点闭合时间,s", str(t_i)],
            ["滤失过程j点压力,MPa", str(P_j)],
            ["滤失过程j点闭合时间,s", str(t_j)],
            ["裂缝延伸系数", str(m)],
            ["裂缝模型", frac_type],
        ]
        df = pd.DataFrame(data, columns=["输入参数", "数值"])

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".1 输入参数")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 创建输入参数表
        table = doc.add_table(df.shape[0] + 1, df.shape[1], style="Light Shading")
        # 插入表头
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        # 插入数据
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        # 设置整个表格字体属性
        table.style.font.size = Pt(12)
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        """模型计算"""
        if frac_type == "PKN":
            beta = math.pi / 4.0
            print(frac_type)
        else:
            beta = 1

        # 确定簇间距和诱导应力系数矩阵
        d = np.zeros((frac_num, frac_num))  # 簇间距
        phi = np.zeros((frac_num, frac_num))  # 诱导应力系数矩阵
        for i in range(frac_num):
            for j in range(frac_num):
                if i != j:
                    d[i][j] = frac_spac * abs(i - j)
                    phi[i][j] = 1 - math.pow(1 + math.pow(Hw / 2 / d[i][j], 2), -1.5)
        console.print("簇间距矩阵：\n", d, style="green")
        console.print("诱导应力系数矩阵：\n", phi, style="green")

        for i in range(frac_num):
            tmp = 0
            for j in range(frac_num):
                tmp += phi[i][j]
            # console.print(tmp, style="green")
            if tmp >= 1:
                console.print("第%d条裂缝未张开" % (i + 1), style="green")
                for k in range(frac_num):
                    phi[k][i] = 0.0000001
                    phi[i][k] = 0.0000001
                    phi[i][i] = 0.999999
                

        console.print("修正后的诱导应力系数矩阵：\n", phi, style="green")

        a = np.zeros((frac_num))  # 缝长系数矩阵
        Pnet = np.zeros((frac_num))  # 停泵时刻各缝净压力矩阵
        for i in range(frac_num):
            tmp = 0
            for j in range(frac_num):
                tmp += phi[i][j]
            a[i] = 1 - tmp
            if a[i] < 0:
                a[i] = 0
            Pnet[i] = a[i] * (ISIP - S1)

        print(ISIP - S1)
        print(a * (ISIP - S1))
        console.print("缝长系数矩阵:\n", a, style="green")
        console.print("停泵时刻各缝净压力矩阵:\n", Pnet, style="green")

        E1 = E / (2 * (1.0 - pow(v, 2)))  # 平面弹性模量
        ti = (t_i - tp) * 1.0 / tp  # 无因次时间
        gi = (4.0 / math.pi) * (
            (1.0 + ti) * math.asin(math.pow((1.0 + ti), -0.5))
            + math.pow(ti, 0.5)
            - math.pi / 2
        )
        tj = (t_j - tp) * 1.0 / tp  # 无因次时间
        gj = (4.0 / math.pi) * (
            (1.0 + tj) * math.asin(math.pow((1.0 + tj), -0.5))
            + math.pow(tj, 0.5)
            - math.pi / 2
        )
        console.print("第i点无因次闭合时间:", gi)
        console.print("第j点无因次闭合时间:", gj)
        b = (16 * beta * math.pow(Hw, 2) / (5 * E1)) * (
            (ISIP - S1)
            - (P_j - P_i)
            / (gj - gi)
            * (np.sqrt(math.pi) * m * math.gamma(m))
            / ((m + 0.5) * math.gamma(m + 0.5)) * eff_c
        )  # 常数分母项
        c = Q * tp  # 常数分子项
        console.print("b, c:", b, c, style="green")
        console.print("常数项:", (c / b), style="green")

        # 联立线性方程组
        # 联立后的缝长系数矩阵
        a_sum = np.zeros((frac_num, frac_num))  # 联立后的缝长系数矩阵
        for i in range(frac_num - 1):
            a_sum[i][i] = 1
            if (a[i] > 0) and (a[i + 1] > 0):
                a_sum[i][i + 1] = 0 - math.pow(a[i] / a[i + 1], nl)

        for i in range(frac_num):
            a_sum[frac_num - 1][i] = a[i]
        console.print("联立后的缝长系数矩阵:\n", a_sum, style="green")

        b_sum = np.zeros((frac_num))  # 联立后的常数项矩阵
        for i in range(frac_num - 1):
            b_sum[i] = 0
        b_sum[frac_num - 1] = c / b
        console.print("联立后的常数项矩阵:\n", b_sum, style="green")

        Lp = linalg.solve(a_sum, b_sum)
        L = Lp * 2

        Wmax = 1.5 * Hw * Pnet / E1
        print(Hw, Pnet, E1)
        C = 0 - 4 * beta * math.pow(Hw, 2) * a / (5 * E1 * Hp * np.sqrt(tp)) * (
            P_j - P_i
        ) / (gj - gi)

        leak_pump = 0  # 泵注过程的总滤失体积
        for i in range(frac_num):
            leak_pump += (
                -16
                * beta
                * math.pow(Hw, 2)
                * (L[i] / 2)
                / (5 * E1)
                * a[i]
                * (P_j - P_i)
                / (gj - gi)
                * np.sqrt(math.pi)
                * m
                * math.gamma(m)
                / ((m + 0.5) * math.gamma(m + 0.5))
            ) * eff_c
            # (4 * np.sqrt(math.pi) * m * math.gamma(m) * C[i] * Hp * L[i] * np.sqrt(tp))/ ((m + 0.5) * math.gamma(m + 0.5))
        leak_shut = 0  # 停泵过程的总滤失体积
        for i in range(frac_num):
            leak_shut += 16 * beta * math.pow(Hw, 2) * (L[i] / 2) / (5 * E1) * Pnet[i]
        V_frac = 0  # 停泵时刻裂缝总体积
        for i in range(frac_num):
            V_frac += 16 * beta * math.pow(Hw, 2) * (L[i] / 2) / (5 * E1) * Pnet[i]
        eff = (Q * tp - leak_pump) / (Q * tp) * 100

        console.print("缝长:\n", L, style="green")
        if frac_num > 1:
            console.print("缝长比例:\n", L[0] / L[1], style="green")
        console.print("泵注过程的总滤失体积:%.3f" % leak_pump, style="green")
        console.print("停泵过程的总滤失体积:%.3f" % leak_shut, style="green")
        console.print("停泵时刻的裂缝总体积:%.3f" % V_frac, style="green")
        console.print("总泵入体积:%.3f" % (Q * tp), style="green")
        console.print("压裂液效率:%.3f%%" % eff, style="green")
        console.print("最大缝宽:\n", Wmax, style="green")
        console.print("滤失系数:\n", C, style="green")

        # 计算滤失面积
        # s = L * 2 * Hw
        # console.print('滤失面积:\n', s, style="green")
        console.print("计算完毕！", style="green")

        table = Table(title="计算参数汇总", show_header=True, header_style="bold magenta")
        table.add_column("参数", style="dim", justify="center")
        table.add_column("数值", style="dim", justify="center")
        table.add_row("缝长", str(L))
        table.add_row("泵注过程的总滤失体积", str(leak_pump))
        table.add_row("停泵过程的总滤失体积/改造体积", str(leak_shut))
        table.add_row("总泵入体积", str(Q * tp))
        table.add_row("压裂液效率", str(eff))
        table.add_row("最大缝宽", str(Wmax))
        table.add_row("滤失系数", str(C))
        # table.add_row('滤失面积', str(s))
        console.print(table)

        data = [
            ["缝长", str(L)],
            ["泵注过程的总滤失体积", str(leak_pump)],
            ["停泵过程的总滤失体积/改造体积", str(leak_shut)],
            ["总泵入体积", str(Q * tp)],
            ["压裂液效率", str(eff)],
            ["最大缝宽", str(Wmax)],
            ["滤失系数", str(C)],
        ]
        df = pd.DataFrame(data, columns=["输出参数", "数值"])

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".2 输出参数")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 创建输入参数表
        table = doc.add_table(df.shape[0] + 1, df.shape[1], style="Light Shading")
        # 插入表头
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        # 插入数据
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        # 设置整个表格字体属性
        table.style.font.size = Pt(12)
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        """可视化"""
        fig = plt.figure()
        ax = Axes3D(fig, auto_add_to_figure=False)
        fig.add_axes(ax)
        for i in range(frac_num):
            Y = np.arange(0, int(L[i] / 2) + 1, 1)
            Z = np.arange(-Hw, Hw + 1, 1)
            Y, Z = np.meshgrid(Y, Z)
            w = Wmax[i] * np.sqrt(np.sqrt(1 - Y / int(L[i] / 2)))
            X = w * np.sqrt(1 - Z * Z / (Hw * Hw))
            # console.print('x:', X)
            # console.print('y:', Y)
            # console.print('w:', w)
            # console.print('z:', Z)
            ax.plot_surface(
                X + i * max(Wmax) * 3, Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, -Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, -Y, Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                X + i * max(Wmax) * 3, -Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
            ax.plot_surface(
                -X + i * max(Wmax) * 3, -Y, -Z, rstride=1, cstride=1, cmap=cm.viridis
            )
        ax.set_xlabel("Fracture Width", rotation=-15)
        ax.set_ylabel("Fracture Length", rotation=50)
        ax.set_zlabel("Fracture Height", rotation=90)

        # 绘制井筒
        # u = np.linspace(0,2*np.pi,50)
        # h = np.linspace(-max(Wmax) * 2,max(Wmax) * 6)
        # y = np.outer(np.sin(u),np.ones(len(h)))
        # z = np.outer(np.cos(u),np.ones(len(h)))
        # x = np.outer(np.ones(len(u)),h)
        # ax.plot_surface(x, y, z, cmap=cm.viridis)

        curr_time = datetime.datetime.now()
        time_str = datetime.datetime.strftime(curr_time, "%Y-%m-%d-%H-%M-%S")
        plt.savefig(file + "\\3D_" + time_str + ".png")

        fig, ax = plt.subplots()
        for i in range(frac_num):
            Y = np.arange(0, int(L[i] / 2) + 1, 0.001)
            w = Wmax[i] * np.sqrt(np.sqrt(1 - Y / int(L[i] / 2)))  # 缝宽作为颜色
            X = np.full((len(Y)), i * frac_spac)

            pos = ax.scatter(X, Y, marker=",", c=w, cmap="RdYlGn_r")
            pos = ax.scatter(X, -Y, marker=",", c=w, cmap="RdYlGn_r")
        
        fig.colorbar(pos, ax=ax)
        # 设置坐标轴范围
        plt.xticks(linspace(-frac_spac / 2, (frac_num - 1) * frac_spac + frac_spac / 2, 5))
        plt.yticks(linspace(-max(Y) - frac_spac / 2, max(Y) + frac_spac / 2, 7))
        plt.xlabel("x")
        plt.ylabel("y")
        plt.savefig(file + "\\2D_" + time_str + ".png")

        # 设置二级标题
        para_heading = doc.add_heading("", level=2)
        para_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_heading.paragraph_format.space_before = Pt(0)
        para_heading.paragraph_format.space_after = Pt(0)
        para_heading.paragraph_format.line_spacing = 1.5
        para_heading.paragraph_format.left_indent = Inches(0)
        para_heading.paragraph_format.right_indent = Inches(0)

        run = para_heading.add_run(str(nn) + ".3 结果可视化")
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(
            file + "\\3D_" + time_str + ".png", width=Cm(12), height=Cm(8)
        )
        paragraph.add_run().add_picture(
            file + "\\2D_" + time_str + ".png", width=Cm(12), height=Cm(8)
        )
        if nn == int(len(data_input.sheet_names)):
            pass
        else:
            doc.add_section()
    doc.save(file + "\\分析报告_" + name + ".docx")
    console.print("分析报告已生成！", style="red")

os.system("pause")